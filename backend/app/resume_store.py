from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
import json
from pathlib import Path
import re

from docx import Document
from pypdf import PdfReader

from app.knowledge import KnowledgeDocument


class ResumeValidationError(ValueError):
    pass


@dataclass(frozen=True)
class ResumeUpload:
    content: bytes
    original_name: str


class ResumeStore:
    def __init__(self, storage_dir: Path, max_upload_bytes: int) -> None:
        self.storage_dir = storage_dir
        self.max_upload_bytes = max_upload_bytes
        self.storage_dir.mkdir(parents=True, exist_ok=True)

    @property
    def metadata_path(self) -> Path:
        return self.storage_dir / "metadata.json"

    @property
    def text_path(self) -> Path:
        return self.storage_dir / "resume.txt"

    def file_path(self, file_type: str) -> Path:
        if file_type not in {"pdf", "docx"}:
            raise ResumeValidationError("Unsupported resume file type.")
        return self.storage_dir / f"Nagendra-Mule-Resume.{file_type}"

    def _validate_size(self, upload: ResumeUpload) -> None:
        if not upload.content:
            raise ResumeValidationError(f"{upload.original_name} is empty.")
        if len(upload.content) > self.max_upload_bytes:
            raise ResumeValidationError(f"{upload.original_name} exceeds the 5 MB upload limit.")

    def _extract_pdf(self, upload: ResumeUpload) -> str:
        self._validate_size(upload)
        if not upload.content.startswith(b"%PDF"):
            raise ResumeValidationError("The selected PDF does not have a valid PDF signature.")
        try:
            reader = PdfReader(BytesIO(upload.content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as exc:
            raise ResumeValidationError("The selected PDF could not be read safely.") from exc

    def _extract_docx(self, upload: ResumeUpload) -> str:
        self._validate_size(upload)
        try:
            document = Document(BytesIO(upload.content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            table_rows = [" | ".join(cell.text for cell in row.cells) for table in document.tables for row in table.rows]
            return "\n".join([*paragraphs, *table_rows])
        except Exception as exc:
            raise ResumeValidationError("The selected DOCX could not be read safely.") from exc

    @staticmethod
    def _normalize_text(text: str) -> str:
        lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
        return "\n".join(line for line in lines if line)

    def save(self, pdf: ResumeUpload | None, docx: ResumeUpload | None) -> dict[str, object]:
        if not pdf and not docx:
            raise ResumeValidationError("Select a PDF, a DOCX, or both.")

        extracted: list[str] = []
        if pdf:
            extracted.append(self._extract_pdf(pdf))
        if docx:
            extracted.append(self._extract_docx(docx))
        text = self._normalize_text(max(extracted, key=len, default=""))
        if len(text) < 80:
            raise ResumeValidationError("The resume does not contain enough extractable text for grounded AI answers.")

        if pdf:
            self.file_path("pdf").write_bytes(pdf.content)
        if docx:
            self.file_path("docx").write_bytes(docx.content)
        self.text_path.write_text(text, encoding="utf-8")

        metadata = {
            "updated_at": datetime.now(timezone.utc).isoformat(),
            "pdf_available": self.file_path("pdf").is_file(),
            "docx_available": self.file_path("docx").is_file(),
            "source_names": [item.original_name for item in (pdf, docx) if item],
        }
        self.metadata_path.write_text(json.dumps(metadata, indent=2), encoding="utf-8")
        return metadata

    def status(self) -> dict[str, object]:
        metadata: dict[str, object] = {}
        if self.metadata_path.is_file():
            try:
                metadata = json.loads(self.metadata_path.read_text(encoding="utf-8"))
            except (json.JSONDecodeError, OSError):
                metadata = {}
        documents = self.load_documents()
        pdf_available = self.file_path("pdf").is_file()
        docx_available = self.file_path("docx").is_file()
        return {
            "available": pdf_available or docx_available,
            "pdf_url": "/api/resume/files/pdf" if pdf_available else None,
            "docx_url": "/api/resume/files/docx" if docx_available else None,
            "updated_at": metadata.get("updated_at"),
            "knowledge_chunks": len(documents),
        }

    def load_documents(self) -> list[KnowledgeDocument]:
        if not self.text_path.is_file():
            return []
        text = self.text_path.read_text(encoding="utf-8").strip()
        if not text:
            return []
        paragraphs = text.splitlines()
        chunks: list[str] = []
        current: list[str] = []
        current_length = 0
        for paragraph in paragraphs:
            if current and current_length + len(paragraph) > 900:
                chunks.append(" ".join(current))
                current = current[-1:]
                current_length = sum(len(item) for item in current)
            current.append(paragraph)
            current_length += len(paragraph)
        if current:
            chunks.append(" ".join(current))
        source_url = "/api/resume/files/pdf" if self.file_path("pdf").is_file() else "/api/resume/files/docx"
        return [
            KnowledgeDocument(
                id=f"uploaded-resume-{index}",
                label="Nagendra Mule Resume",
                section=f"Resume section {index}",
                text=chunk,
                url=source_url,
                metadata={"file_category": "resume", "uploaded": True},
            )
            for index, chunk in enumerate(chunks, start=1)
        ]
